"""Генерация Word: описание модульных и интеграционных тестов cms-workflow-service."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).resolve().parent.parent / "testing-documentation.docx"

CODE_FONT = "Consolas"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_snippet(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F5F5F5")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for line in code.rstrip().split("\n"):
        run = paragraph.add_run(line + "\n")
        run.font.name = CODE_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def add_test_case(
    doc: Document,
    index: int,
    method_name: str,
    description: str,
    code: str,
) -> None:
    doc.add_heading(f"{index}. {method_name}", level=4)
    doc.add_paragraph(description).paragraph_format.space_after = Pt(6)
    doc.add_paragraph("Фрагмент кода теста:").runs[0].italic = True
    add_code_snippet(doc, code)


# (класс, файл, краткое описание класса, [(имя it, описание, код), ...])
MODULAR_TESTS: list[tuple[str, str, str, list[tuple[str, str, str]]]] = [
    (
        "IngestIncidentTopicUseCase",
        "src/core/incident-management/use-cases/ingest-incident-topic.use-case.spec.ts",
        "Use-case приёма сообщений из Kafka-топика инцидентов: валидирует payload "
        "и ставит задачу в outbox для последующей обработки IncidentResolverService.",
        [
            (
                "ignores message when companyId is missing",
                "Проверяет, что при пустом companyId (невалидный payload) метод execute "
                "завершается без записи в outbox — сообщение игнорируется.",
                """  it('ignores message when companyId is missing', async () => {
    await useCase.execute({
      companyId: '',
      integrationId: 1,
      riskObjectId: 'risk-1',
      rules: [],
    } as IncidentTopicMessage);

    const pending = await outboxRepository.getPending(10);
    expect(pending).toHaveLength(0);
  });""",
            ),
            (
                "persists pending outbox message for valid incident topic",
                "Проверяет, что при корректном IncidentTopicMessage создаётся одна запись "
                "outbox со статусом pending, топиком incident_topic.received и payload, "
                "содержащим companyId, integrationId, riskObjectId, documentId и rules.",
                """  it('persists pending outbox message for valid incident topic', async () => {
    const message: IncidentTopicMessage = {
      companyId: 'company-1',
      integrationId: 42,
      riskObjectId: 'risk-1',
      documentId: 'doc-1',
      rules: [
        {
          rulesId: 'rule-1',
          rulePriority: 'HIGH',
          responsible_user_id: 'user-1',
          result: 'fail',
          found: true,
          details: { code: 'X1' },
        },
      ],
    };

    await useCase.execute(message);

    const pending = await outboxRepository.getPending(10);
    expect(pending).toHaveLength(1);
    expect(pending[0]).toMatchObject({
      topic: 'incident_topic.received',
      status: 'pending',
      payload: expect.objectContaining({
        companyId: 'company-1',
        integrationId: 42,
        riskObjectId: 'risk-1',
        documentId: 'doc-1',
        rules: message.rules,
      }),
    });
  });""",
            ),
        ],
    ),
    (
        "OutboxProcessorService",
        "src/core/outbox/services/outbox-processor.service.spec.ts",
        "Фоновый планировщик обработки outbox: читает pending-сообщения, "
        "маршрутизирует по topic и помечает обработанные (с очисткой).",
        [
            (
                "marks non-incident topics as processed without calling resolver",
                "Проверяет, что для топика, отличного от incident_topic.received, "
                "IncidentResolverService.resolveOutboxMessage не вызывается, "
                "а очередь pending после processPendingMessages пуста (сообщение обработано и удалено).",
                """  it('marks non-incident topics as processed without calling resolver', async () => {
    await seedMessage({ id: 'msg-1', topic: 'other.topic' });

    await service.processPendingMessages();

    expect(incidentResolver.resolveOutboxMessage).not.toHaveBeenCalled();
    const pending = await outboxRepository.getPending(10);
    expect(pending).toHaveLength(0);
  });""",
            ),
            (
                "delegates incident_topic.received messages to incident resolver",
                "Проверяет, что сообщение с топиком incident_topic.received передаётся "
                "в resolveOutboxMessage ровно один раз с ожидаемым id и topic, "
                "после чего pending-очередь пуста.",
                """  it('delegates incident_topic.received messages to incident resolver', async () => {
    await seedMessage({
      id: 'msg-2',
      topic: 'incident_topic.received',
      payload: { companyId: 'c1', integrationId: 1, riskObjectId: 'r1', rules: [] },
    });

    await service.processPendingMessages();

    expect(incidentResolver.resolveOutboxMessage).toHaveBeenCalledTimes(1);
    expect(incidentResolver.resolveOutboxMessage).toHaveBeenCalledWith(
      expect.objectContaining({ id: 'msg-2', topic: 'incident_topic.received' }),
    );
    const pending = await outboxRepository.getPending(10);
    expect(pending).toHaveLength(0);
  });""",
            ),
        ],
    ),
    (
        "InMemoryOutboxRepository",
        "src/infrastructure/outbox/persistence/in-memory-outbox.repository.spec.ts",
        "In-memory реализация OutboxRepositoryPort для тестов и локальной разработки.",
        [
            (
                "returns only pending messages up to the requested limit",
                "Проверяет метод getPending(limit): возвращаются только сообщения "
                "со статусом pending, processed не попадают в выборку; "
                "соблюдается ограничение limit (в тесте limit=1 при двух pending).",
                """  it('returns only pending messages up to the requested limit', async () => {
    const repository = new InMemoryOutboxRepository();

    await repository.add({
      id: 'pending-1',
      topic: 'incident_topic.received',
      payload: {},
      createdAt: new Date(),
      status: 'pending',
    });
    await repository.add({
      id: 'pending-2',
      topic: 'incident_topic.received',
      payload: {},
      createdAt: new Date(),
      status: 'pending',
    });
    await repository.add({
      id: 'processed-1',
      topic: 'incident_topic.received',
      payload: {},
      createdAt: new Date(),
      status: 'processed',
    });

    const pending = await repository.getPending(1);

    expect(pending).toHaveLength(1);
    expect(pending[0].status).toBe('pending');
  });""",
            ),
        ],
    ),
]

INTEGRATION_TESTS: list[tuple[str, str, str, list[tuple[str, str, str]]]] = [
    (
        "RejectCaseUseCase",
        "test/business-logic.integration-spec.ts",
        "Сценарий отклонения кейса расследования: смена статуса кейса, "
        "создание action plan с комментарием, пересчёт статуса инцидента. "
        "Тестируется в связке с in-memory репозиториями через Nest TestingModule.",
        [
            (
                "rejects investigating case and resolves incident when it is the only case",
                "Проверяет успешный execute: кейс INVESTIGATING → REJECTED, "
                "единственный кейс инцидента → инцидент RESOLVED с resolvedDate, "
                "создаётся action plan с переданным comment.",
                """  it('rejects investigating case and resolves incident when it is the only case', async () => {
    const { caseId, incidentId } = await seedIncidentGraph();

    const result = await rejectCaseUseCase.execute(caseId, {
      comment: 'Not enough evidence',
    });

    expect(result).toEqual({
      caseId,
      caseStatus: 'REJECTED',
      incidentId,
      incidentStatus: 'RESOLVED',
    });

    const savedIncident = await incidentRepository.findOne({
      where: { id: incidentId },
    });
    expect(savedIncident?.status).toBe('RESOLVED');
    expect(savedIncident?.resolvedDate).toBeInstanceOf(Date);

    const actionPlans = actionPlanRepository.getAll();
    expect(actionPlans).toHaveLength(1);
    expect(actionPlans[0]).toMatchObject({
      caseId,
      incidentId,
      comment: 'Not enough evidence',
    });
  });""",
            ),
            (
                "requires comment when rejecting a case",
                "Проверяет валидацию: при отсутствии comment в payload "
                "выбрасывается BadRequestException.",
                """  it('requires comment when rejecting a case', async () => {
    const { caseId } = await seedIncidentGraph();

    await expect(rejectCaseUseCase.execute(caseId, {})).rejects.toBeInstanceOf(
      BadRequestException,
    );
  });""",
            ),
            (
                "allows rejection only from INVESTIGATING status",
                "Проверяет бизнес-правило: отклонение из статуса ASSIGNED запрещено — "
                "ожидается BadRequestException.",
                """  it('allows rejection only from INVESTIGATING status', async () => {
    const { caseId } = await seedIncidentGraph({ caseStatus: 'ASSIGNED' });

    await expect(
      rejectCaseUseCase.execute(caseId, { comment: 'Too early' }),
    ).rejects.toBeInstanceOf(BadRequestException);
  });""",
            ),
        ],
    ),
    (
        "ReopenCaseUseCase",
        "test/business-logic.integration-spec.ts",
        "Сценарий повторного открытия отклонённого кейса (REJECTED → INVESTIGATING).",
        [
            (
                "reopens case from REJECTED back to INVESTIGATING",
                "Проверяет успешный execute: возвращается caseStatus INVESTIGATING, "
                "запись в репозитории обновлена.",
                """  it('reopens case from REJECTED back to INVESTIGATING', async () => {
    const { caseId } = await seedIncidentGraph({ caseStatus: 'REJECTED' });

    const result = await reopenCaseUseCase.execute(caseId);

    expect(result).toEqual({
      caseId,
      caseStatus: 'INVESTIGATING',
    });

    const savedCase = await caseRepository.findOne({ where: { id: caseId } });
    expect(savedCase?.status).toBe('INVESTIGATING');
  });""",
            ),
            (
                "allows reopen only from REJECTED status",
                "Проверяет бизнес-правило: reopen из INVESTIGATING запрещён — "
                "ожидается BadRequestException.",
                """  it('allows reopen only from REJECTED status', async () => {
    const { caseId } = await seedIncidentGraph({ caseStatus: 'INVESTIGATING' });

    await expect(reopenCaseUseCase.execute(caseId)).rejects.toBeInstanceOf(
      BadRequestException,
    );
  });""",
            ),
        ],
    ),
    (
        "AssignIncidentToMeUseCase",
        "test/business-logic.integration-spec.ts",
        "Сценарий назначения инцидента на себя (claim): для каждой доступной находки "
        "создаётся кейс (case) со статусом ASSIGNED, обновляется finding.assignedUserId "
        "и статус инцидента. Используются in-memory репозитории, DataSource-транзакция "
        "и мок HTTP Auth (fetch).",
        [
            (
                "creates ASSIGNED case when manager claims unassigned finding",
                "Проверяет создание case: для инцидента с незакреплённой находкой "
                "execute возвращает один кейс со статусом ASSIGNED, assignedUserId менеджера, "
                "finding получает assignedUserId, инцидент переходит в IN_PROGRESS, "
                "в репозитории cases ровно одна запись.",
                """  it('creates ASSIGNED case when manager claims unassigned finding', async () => {
    const incidentId = randomUUID();
    const findingId = randomUUID();

    incidentRepository.seed({
      id: incidentId,
      companyId: 'company-1',
      integrationId: 1,
      riskObjectId: 'risk-1',
      departmentId: null,
      documentId: null,
      status: 'OPEN',
      resolvedDate: null,
    } as IncidentOrmEntity);

    findingRepository.seed({
      id: findingId,
      incidentId,
      priority: 'HIGH',
      assignedUserId: null,
      rulesId: null,
      detectedAt: new Date(),
      deadline: null,
      details: { code: 'R1' },
    } as FindingOrmEntity);

    const cases = await assignIncidentToMeUseCase.execute(incidentId);

    expect(cases).toHaveLength(1);
    expect(cases[0]).toMatchObject({
      incidentId,
      findingId,
      assignedUserId: 'user-1',
      status: 'ASSIGNED',
    });
    expect(cases[0].id).toBeDefined();

    const updatedFinding = await findingRepository.findOne({
      where: { id: findingId },
    });
    expect(updatedFinding?.assignedUserId).toBe('user-1');

    const updatedIncident = await incidentRepository.findOne({
      where: { id: incidentId },
    });
    expect(updatedIncident?.status).toBe('IN_PROGRESS');

    expect(caseRepository.getAll()).toHaveLength(1);
  });""",
            ),
        ],
    ),
    (
        "IncidentResolverService",
        "test/business-logic.integration-spec.ts",
        "Создание инцидента и находок (findings) из сообщения outbox после валидации payload. "
        "Мок HTTP Monitoring Service для departmentId. In-memory DataSource и репозитории.",
        [
            (
                "creates OPEN incident and findings from outbox message",
                "Проверяет resolveOutboxMessage: в БД появляется инцидент OPEN с companyId, "
                "integrationId, riskObjectId, departmentId; одна finding с priority, rulesId, "
                "assignedUserId и details из правила.",
                """  it('creates OPEN incident and findings from outbox message', async () => {
    await incidentResolverService.resolveOutboxMessage({
      id: randomUUID(),
      topic: 'incident_topic.received',
      payload: {
        companyId: 'company-1',
        integrationId: 99,
        riskObjectId: 'risk-99',
        documentId: 'doc-1',
        rules: [
          {
            rulesId: 'rule-1',
            rulePriority: 'HIGH',
            responsible_user_id: 'user-2',
            result: 'fail',
            found: true,
            details: { code: 'R1' },
          },
        ],
      },
      createdAt: new Date(),
      status: 'pending',
    });

    const incidents = incidentRepository.getAll();
    expect(incidents).toHaveLength(1);
    expect(incidents[0]).toMatchObject({
      companyId: 'company-1',
      integrationId: 99,
      riskObjectId: 'risk-99',
      documentId: 'doc-1',
      status: 'OPEN',
      departmentId: 'dept-1',
    });

    const findings = findingRepository.getAll();
    expect(findings).toHaveLength(1);
    expect(findings[0]).toMatchObject({
      incidentId: incidents[0].id,
      priority: 'HIGH',
      assignedUserId: 'user-2',
      rulesId: 'rule-1',
    });
  });""",
            ),
        ],
    ),
    (
        "CreateActionPlanUseCase",
        "test/business-logic.integration-spec.ts",
        "Создание плана корректирующих действий с задачами для кейса: action plan, tasks, "
        "перевод кейса в статус ACTION_PLAN. Транзакция через in-memory DataSource.",
        [
            (
                "creates action plan with tasks and sets case to ACTION_PLAN",
                "Проверяет execute: возвращается план с title, description, tasks (TODO), "
                "caseStatus ACTION_PLAN; в репозиториях action_plans и action_plan_tasks "
                "по одной записи; статус кейса обновлён.",
                """  it('creates action plan with tasks and sets case to ACTION_PLAN', async () => {
    const { caseId, incidentId } = await seedIncidentGraph({
      caseStatus: 'INVESTIGATING',
    });

    const result = await createActionPlanUseCase.execute({
      caseId,
      title: 'Corrective plan',
      description: 'Steps to fix the issue',
      tasks: [
        {
          title: 'Task 1',
          description: 'Do something',
          priority: 'HIGH',
          dueDate: '2026-12-31T00:00:00.000Z',
        },
      ],
    });

    expect(result).toMatchObject({
      caseId,
      caseStatus: 'ACTION_PLAN',
      title: 'Corrective plan',
    });
    expect(result.tasks).toHaveLength(1);

    expect(actionPlanRepository.getAll()).toHaveLength(1);
    expect(actionPlanTaskRepository.getAll()).toHaveLength(1);

    const savedCase = await caseRepository.findOne({ where: { id: caseId } });
    expect(savedCase?.status).toBe('ACTION_PLAN');
  });""",
            ),
        ],
    ),
    (
        "IngestIncidentTopicUseCase (интеграция с DI)",
        "test/business-logic.integration-spec.ts",
        "Проверка сквозной связки Nest DI: use-case + InMemoryOutboxRepository "
        "через токен OUTBOX_REPOSITORY.",
        [
            (
                "ingests incident topic into outbox through Nest DI wiring",
                "Проверяет, что после execute в outbox появляется pending-сообщение "
                "с корректным payload (companyId, integrationId, riskObjectId).",
                """  it('ingests incident topic into outbox through Nest DI wiring', async () => {
    await ingestIncidentTopicUseCase.execute({
      companyId: 'company-1',
      integrationId: 7,
      riskObjectId: 'risk-42',
      rules: [
        {
          rulesId: 'rule-1',
          rulePriority: 'LOW',
          responsible_user_id: null,
          result: 'fail',
          found: true,
          details: {},
        },
      ],
    });

    const pending = await outboxRepository.getPending(5);
    expect(pending).toHaveLength(1);
    expect(pending[0].payload).toEqual(
      expect.objectContaining({
        companyId: 'company-1',
        integrationId: 7,
        riskObjectId: 'risk-42',
      }),
    );
  });""",
            ),
        ],
    ),
]


def add_part(
    doc: Document,
    part_title: str,
    part_intro: str,
    classes: list[tuple[str, str, str, list[tuple[str, str, str]]]],
    run_commands: str,
) -> None:
    doc.add_heading(part_title, level=1)
    doc.add_paragraph(part_intro).paragraph_format.space_after = Pt(6)
    doc.add_paragraph(f"Запуск: {run_commands}.").paragraph_format.space_after = Pt(12)

    test_index = 1
    for class_idx, (class_name, file_path, class_desc, tests) in enumerate(classes, start=1):
        doc.add_heading(f"{class_idx}. {class_name}", level=2)
        doc.add_paragraph(f"Файл: {file_path}.").paragraph_format.space_after = Pt(4)
        doc.add_paragraph(class_desc).paragraph_format.space_after = Pt(8)

        if "IngestIncidentTopicUseCase" in class_name and "интеграция" not in class_name:
            tested = "execute(message: IncidentTopicMessage)"
        elif "IngestIncidentTopicUseCase" in class_name:
            tested = "execute(message) — через Nest DI и OUTBOX_REPOSITORY"
        elif class_name == "OutboxProcessorService":
            tested = "processPendingMessages()"
        elif class_name == "InMemoryOutboxRepository":
            tested = "add(), getPending(limit)"
        elif class_name == "RejectCaseUseCase":
            tested = "execute(caseId, payload)"
        elif class_name == "ReopenCaseUseCase":
            tested = "execute(caseId)"
        elif class_name == "AssignIncidentToMeUseCase":
            tested = "execute(incidentId)"
        elif class_name == "IncidentResolverService":
            tested = "resolveOutboxMessage(message)"
        elif class_name == "CreateActionPlanUseCase":
            tested = "execute(payload)"
        else:
            tested = "—"

        p = doc.add_paragraph()
        p.add_run("Тестируемые методы: ").bold = True
        p.add_run(tested)

        for method_name, description, code in tests:
            add_test_case(doc, test_index, method_name, description, code)
            test_index += 1

        doc.add_paragraph()


def main() -> None:
    doc = Document()

    title = doc.add_heading(
        "Документация по тестированию cms-workflow-service",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Описание модульных и интеграционных тестов бизнес-логики микросервиса "
        "TrustFlow CMS Workflow. Фреймворк: Jest; для интеграционных сценариев — "
        "@nestjs/testing и in-memory репозитории."
    ).paragraph_format.space_after = Pt(12)

    summary = doc.add_table(rows=3, cols=2)
    summary.style = "Table Grid"
    rows_data = (
        ("Модульные тесты", "5"),
        ("Интеграционные тесты", "9"),
        ("Всего тестовых методов (it)", "14"),
    )
    for i, (label, value) in enumerate(rows_data):
        summary.rows[i].cells[0].text = label
        summary.rows[i].cells[1].text = value
        for run in summary.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    add_part(
        doc,
        "1. Модульное тестирование",
        "Изолированная проверка отдельных классов: зависимости подменяются моками "
        "или in-memory реализациями. Не требуется PostgreSQL и внешние HTTP-сервисы.",
        MODULAR_TESTS,
        "npm test",
    )

    doc.add_page_break()

    add_part(
        doc,
        "2. Интеграционное тестирование",
        "Проверка связки use-case + persistence + Nest DI. Репозитории TypeORM "
        "заменены на InMemoryTypeOrmRepository; outbox — InMemoryOutboxRepository.",
        INTEGRATION_TESTS,
        "npm run test:integration",
    )

    output = Path(os.environ.get("DOCX_OUTPUT", OUTPUT))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Written: {output}")


if __name__ == "__main__":
    main()
