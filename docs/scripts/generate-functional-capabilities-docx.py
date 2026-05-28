"""Генерация Word: функциональные возможности cms-workflow-service по ролям."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = Path(__file__).resolve().parent.parent / "functional-capabilities.docx"

# (№, функция, метод/путь, роли, примечание)
SECTIONS: list[tuple[str, str, list[tuple[str, str, str, str, str]]]] = [
    (
        "Фоновые и интеграционные возможности",
        "Выполняются сервисом без участия пользователя CMS (не REST API).",
        [
            (
                "1",
                "Приём событий о срабатывании правил из Kafka (топик KAFKA_INCIDENT_TOPIC)",
                "Kafka consumer",
                "—",
                "Сохранение в outbox_messages",
            ),
            (
                "2",
                "Создание инцидента и находок из outbox (IncidentResolverService)",
                "Фоновый обработчик outbox",
                "—",
                "Статус инцидента OPEN",
            ),
            (
                "3",
                "Повторная обработка и ретраи сообщений outbox",
                "OutboxProcessorService",
                "—",
                "Статусы pending / processed / failed",
            ),
        ],
    ),
    (
        "Инциденты (REST: /api/incidents)",
        "Роли из CMS Auth. Для SUPERVISOR при отсутствии employeeId в профиле нужен заголовок EmployeeId.",
        [
            (
                "4",
                "Список «моих» инцидентов с находками и кейсами",
                "GET /api/incidents/my",
                "MANAGER, SUPERVISOR",
                "MANAGER — свои назначения; SUPERVISOR — подчинённые отдела. EXECUTIVE — пустой список",
            ),
            (
                "5",
                "Статистика по инцидентам (новые, решённые, ожидают верификации и др.)",
                "GET /api/incidents/my/stats",
                "MANAGER, SUPERVISOR, EXECUTIVE",
                "EXECUTIVE — вся компания; SUPERVISOR — отдел; MANAGER — свои назначения",
            ),
            (
                "6",
                "KPI по менеджерам с назначениями",
                "GET /api/incidents/kpi/managers",
                "SUPERVISOR, EXECUTIVE",
                "SUPERVISOR — подчинённые; EXECUTIVE — все по компании",
            ),
            (
                "7",
                "Операционная сводка (конвейер, «залипшие» >14 дней, топ объектов риска)",
                "GET /api/incidents/overview",
                "SUPERVISOR, EXECUTIVE",
                "",
            ),
            (
                "8",
                "Эффективность правил (кейсы REJECTED / CLOSED по rulesId)",
                "GET /api/incidents/rule-effectiveness",
                "SUPERVISOR, EXECUTIVE",
                "",
            ),
            (
                "9",
                "Постраничный список полных отчётов по инцидентам",
                "GET /api/incidents/reports",
                "SUPERVISOR, EXECUTIVE",
                "Фильтры: page, limit, incidentId, documentId, status",
            ),
            (
                "10",
                "Проблемные зоны (повторные срабатывания по documentId за месяц)",
                "GET /api/incidents/problem-areas",
                "SUPERVISOR, EXECUTIVE",
                "Query: month (YYYY-MM)",
            ),
            (
                "11",
                "Полный отчёт по инциденту (кейсы, планы, задачи, верификации)",
                "GET /api/incidents/:incidentId/report",
                "Авторизованный пользователь",
                "Явная проверка роли в use case отсутствует",
            ),
            (
                "12",
                "Краткое представление инцидента (findings, documentId, интеграция)",
                "GET /api/incidents/:incidentId/view",
                "Авторизованный пользователь",
                "Явная проверка роли в use case отсутствует",
            ),
            (
                "13",
                "Назначить инцидент себе: создание кейсов по находкам",
                "POST /api/incidents/:incidentId/assign-to-me",
                "MANAGER (типично)",
                "Пользователь той же companyId; свои или неназначенные findings",
            ),
        ],
    ),
    (
        "Кейсы (REST: /api/cases, /api/v1/cases)",
        "Коллаборация: SUPERVISOR — любой кейс; иначе — исполнитель кейса или другой кейс того же инцидента с тем же исполнителем.",
        [
            (
                "14",
                "Список всех кейсов в системе",
                "GET /api/cases",
                "Авторизованный пользователь",
                "Без фильтрации по роли в use case",
            ),
            (
                "15",
                "Список своих кейсов (ruleId, priority, status, deadline)",
                "GET /api/v1/cases/my",
                "MANAGER",
                "",
            ),
            (
                "16",
                "Статистика по своим кейсам",
                "GET /api/v1/cases/my/stats",
                "MANAGER",
                "Только assignedUserId = текущий пользователь",
            ),
            (
                "17",
                "Карточка кейса (правило, условие, details находки)",
                "GET /api/v1/cases/:caseId/view",
                "MANAGER",
                "",
            ),
            (
                "18",
                "Обновление материалов расследования",
                "PATCH /api/cases/:caseId/investigation",
                "MANAGER",
                "Только исполнитель кейса (assignedUserId); статус → INVESTIGATING",
            ),
            (
                "19",
                "Отклонение кейса с указанием причины",
                "PUT /api/cases/:caseId/reject",
                "MANAGER (типично)",
                "Только из INVESTIGATING; пересчёт статуса инцидента",
            ),
            (
                "20",
                "Повторное открытие отклонённого кейса",
                "PUT /api/cases/:caseId/reopen",
                "MANAGER (типично)",
                "Только из REJECTED → INVESTIGATING",
            ),
            (
                "21",
                "Просмотр комментариев кейса",
                "GET /api/cases/:caseId/comments",
                "MANAGER, SUPERVISOR",
                "Правила коллаборации по кейсу",
            ),
            (
                "22",
                "Добавление комментария",
                "POST /api/cases/:caseId/comments, POST /api/v1/cases/:caseId/comments",
                "MANAGER, SUPERVISOR",
                "Правила коллаборации",
            ),
            (
                "23",
                "Список вложений кейса",
                "GET /api/cases/:caseId/attachments",
                "MANAGER, SUPERVISOR",
                "Правила коллаборации",
            ),
            (
                "24",
                "Загрузка вложения (MinIO)",
                "POST /api/cases/:caseId/attachments, POST /api/v1/cases/:caseId/attachments",
                "MANAGER, SUPERVISOR",
                "Правила коллаборации",
            ),
            (
                "25",
                "Скачивание вложения",
                "GET /api/cases/:caseId/attachments/:id/download",
                "MANAGER, SUPERVISOR",
                "Правила коллаборации",
            ),
            (
                "26",
                "Удаление вложения",
                "DELETE /api/cases/:caseId/attachments/:id",
                "MANAGER, SUPERVISOR",
                "Кейс в статусе INVESTIGATING; коллаборация",
            ),
        ],
    ),
    (
        "Расследования (REST: /api/investigations)",
        "",
        [
            (
                "27",
                "Список всех записей расследований",
                "GET /api/investigations",
                "Авторизованный пользователь",
                "Без проверки роли в use case",
            ),
        ],
    ),
    (
        "Планы корректирующих действий (REST: /api/action-plans)",
        "",
        [
            (
                "28",
                "Список планов по своим кейсам",
                "GET /api/action-plans",
                "MANAGER",
                "Иначе пустой список",
            ),
            (
                "29",
                "Создание или дополнение плана и задач",
                "POST /api/action-plans",
                "MANAGER (типично)",
                "Явная проверка роли отсутствует; кейс → ACTION_PLAN",
            ),
            (
                "30",
                "Обновление полей плана (title, description, comment)",
                "PATCH /api/action-plans/:planId",
                "MANAGER",
                "Коллаборация по кейсу",
            ),
            (
                "31",
                "Отправка плана на верификацию",
                "POST /api/action-plans/:planId/submit",
                "MANAGER",
                "Назначение верификатора из company-info; кейс → WAITING_VERIFICATION",
            ),
            (
                "32",
                "Подтверждение (приёмка) плана",
                "POST /api/action-plans/:planId/confirm",
                "SUPERVISOR, EXECUTIVE",
                "approved=true, comments обязательны",
            ),
            (
                "33",
                "Возврат плана на доработку",
                "POST /api/action-plans/:planId/return-for-revision",
                "SUPERVISOR, EXECUTIVE",
                "Кейс → ACTION_PLAN",
            ),
            (
                "34",
                "Удаление задачи из плана",
                "DELETE /api/action-plans/:planId/tasks/:taskId",
                "MANAGER, руководитель отдела",
                "Исполнитель кейса или department manager / верификатор",
            ),
            (
                "35",
                "Список файлов-доказательств по задаче",
                "GET .../tasks/:taskId/evidences",
                "MANAGER, SUPERVISOR",
                "Кейс ACTION_IN_PROGRESS; коллаборация",
            ),
            (
                "36",
                "Загрузка доказательства по задаче",
                "POST .../tasks/:taskId/evidences",
                "MANAGER, SUPERVISOR",
                "Кейс ACTION_IN_PROGRESS",
            ),
            (
                "37",
                "Скачивание доказательства",
                "GET .../evidences/:evidenceId/download",
                "MANAGER, SUPERVISOR",
                "Кейс ACTION_IN_PROGRESS",
            ),
        ],
    ),
    (
        "Верификация планов (REST: /api/supervisor/verification)",
        "",
        [
            (
                "38",
                "Очередь планов, ожидающих верификации",
                "GET /api/supervisor/verification/pending",
                "SUPERVISOR, EXECUTIVE",
                "SUPERVISOR — где назначен верификатором; EXECUTIVE — по компании",
            ),
            (
                "39",
                "Подтверждение верификации плана",
                "PUT /api/supervisor/verification/:actionPlanId",
                "SUPERVISOR, EXECUTIVE",
                "Аналог POST .../confirm",
            ),
        ],
    ),
    (
        "Задачи плана (REST: /api/tasks)",
        "Видимость задач после приёмки плана (showTasks=true, кейс ACTION_IN_PROGRESS).",
        [
            (
                "40",
                "Список задач текущего пользователя",
                "GET /api/tasks/my",
                "MANAGER",
                "По кейсам, где пользователь assignedUserId; showTasks=true",
            ),
            (
                "41",
                "Статистика по задачам",
                "GET /api/tasks/my/stats",
                "MANAGER",
                "",
            ),
            (
                "42",
                "Просмотр задачи по id",
                "GET /api/tasks/:taskId",
                "MANAGER, SUPERVISOR",
                "Коллаборация; showTasks=true",
            ),
            (
                "43",
                "Обновление статуса и описания прогресса задачи",
                "PATCH /api/tasks/:taskId",
                "MANAGER, SUPERVISOR",
                "TODO / IN_PROGRESS / DONE; при всех DONE — кейс CLOSED",
            ),
            (
                "44",
                "Завершение задачи с описанием доказательства",
                "POST /api/tasks/:taskId/complete",
                "MANAGER, SUPERVISOR",
                "Статус DONE; может закрыть кейс и инцидент",
            ),
            (
                "45",
                "Загрузка доказательства (альтернативный endpoint)",
                "POST /api/tasks/:taskId/evidence",
                "MANAGER, SUPERVISOR",
                "Кейс ACTION_IN_PROGRESS",
            ),
        ],
    ),
    (
        "Служебные",
        "",
        [
            (
                "46",
                "Проверка доступности сервиса",
                "GET /",
                "Без роли",
                "Health-check",
            ),
            (
                "47",
                "Интерактивная документация OpenAPI",
                "GET /api/docs",
                "Без роли (в dev)",
                "Swagger UI",
            ),
        ],
    ),
]

ROLES_LEGEND = [
    ("MANAGER", "Менеджер / исполнитель: ведёт кейсы, расследование, план и задачи по своим назначениям."),
    ("SUPERVISOR", "Руководитель подразделения: обзор подчинённых, верификация планов, расширенная коллаборация по кейсам."),
    ("EXECUTIVE", "Руководитель компании: аналитика и отчёты по всей компании, верификация планов."),
    (
        "Авторизованный пользователь",
        "Любой пользователь с валидным Authorization; отдельная проверка роли в коде не выполняется.",
    ),
]

# Варианты использования: (наименование, [(запрос, описание, исключения), ...])
USE_CASE_API: list[tuple[str, list[tuple[str, str, str]]]] = [
    (
        "Просмотреть кейс",
        [
            (
                "GET /api/v1/cases/my",
                "MANAGER. Заголовок Authorization.\n"
                "Ответ: список своих кейсов (правило, приоритет, статус, срок).",
                "401 — нет доступа / не MANAGER.\n404 — не применяется.",
            ),
            (
                "GET /api/v1/cases/{caseId}/view",
                "MANAGER. Path: caseId.\n"
                "Ответ: карточка кейса — правило, детали срабатывания, данные расследования.",
                "401, 404 — нет доступа или кейс не найден.",
            ),
            (
                "GET /api/action-plans",
                "MANAGER.\n"
                "Ответ: планы и задачи по своим кейсам (пустой список для других ролей).",
                "401 — ошибка авторизации.",
            ),
        ],
    ),
    (
        "Создать план корректирующих действий",
        [
            (
                "POST /api/action-plans",
                "JSON: caseId, title, description; опционально список задач (срок, приоритет).\n"
                "Ответ: созданный план, статус кейса ACTION_PLAN; при повторе — дополнение задач.",
                "400 — неполные или неверные данные.\n404 — кейс не найден.\n401 — нет авторизации.",
            ),
            (
                "PATCH /api/action-plans/{planId}",
                "MANAGER. JSON: title, description или comment.\n"
                "Ответ: обновлённый план.",
                "403 — нет прав.\n404 — план не найден.",
            ),
            (
                "POST /api/action-plans/{planId}/submit",
                "MANAGER. Заголовок EmployeeId.\n"
                "Ответ: план отправлен на проверку, кейс WAITING_VERIFICATION.",
                "403 — не менеджер.\n400 — нет EmployeeId или руководителя в company-info.\n404 — план не найден.",
            ),
        ],
    ),
    (
        "Провести расследование",
        [
            (
                "PATCH /api/cases/{caseId}/investigation",
                "MANAGER (исполнитель кейса). JSON: заметки, причина, нужны ли корректирующие действия.\n"
                "Ответ: кейс INVESTIGATING и сохранённое расследование.",
                "403 — не ваш кейс / другая компания.\n404 — кейс не найден.\n401 — нет авторизации.",
            ),
            (
                "GET /api/cases/{caseId}/comments",
                "MANAGER или SUPERVISOR.\n"
                "Ответ: лента комментариев.",
                "403, 404 — нет доступа или кейс не найден.",
            ),
            (
                "POST /api/cases/{caseId}/attachments",
                "Файл в multipart (поле file).\n"
                "Ответ: метаданные вложения в MinIO.",
                "400 — файл не передан.\n403, 404 — нет доступа.",
            ),
            (
                "PUT /api/cases/{caseId}/reject",
                "JSON: comment (причина отклонения).\n"
                "Ответ: кейс REJECTED, пересчитан статус инцидента.",
                "400 — нет комментария, неверный статус или уже есть план.\n404 — кейс не найден.",
            ),
        ],
    ),
    (
        "Выполнить план корректирующих действий",
        [
            (
                "POST /api/action-plans/{planId}/confirm",
                "SUPERVISOR / EXECUTIVE. JSON: подтверждение и комментарий.\n"
                "Ответ: план принят, кейс ACTION_IN_PROGRESS, задачи видны исполнителю.",
                "403 — нет роли.\n400 — нет комментария.\n404 — план не найден.",
            ),
            (
                "GET /api/tasks/my",
                "Исполнитель кейса, план уже принят.\n"
                "Ответ: список своих задач плана.",
                "200 [] — задач нет.\n401, 403 — нет доступа.",
            ),
            (
                "PATCH /api/tasks/{taskId}",
                "JSON: статус задачи и/или описание прогресса.\n"
                "Ответ: обновлённая задача; при завершении всех задач — закрытие кейса и инцидента.",
                "400 — пустое тело или неверный статус.\n403, 404 — нет доступа или задача не найдена.",
            ),
            (
                "POST /api/tasks/{taskId}/complete",
                "JSON: описание выполненной работы.\n"
                "Ответ: задача завершена.",
                "400 — нет описания.\n403, 404 — как у PATCH.",
            ),
            (
                "POST /api/tasks/{taskId}/evidence",
                "Кейс в стадии исполнения. Файл в multipart.\n"
                "Ответ: ссылка на файл-доказательство.",
                "403 — кейс не в ACTION_IN_PROGRESS.\n400 — файл не передан.",
            ),
        ],
    ),
    (
        "Просмотреть Dashboard и KPI",
        [
            (
                "GET /api/incidents/my/stats",
                "EXECUTIVE / SUPERVISOR / MANAGER (разный охват данных). EmployeeId для руководителя при необходимости.\n"
                "Ответ: сводка по инцидентам — количество, статусы, сроки, просрочки, верификации.",
                "401, 403 — нет доступа.\n400 — нет EmployeeId у SUPERVISOR.",
            ),
            (
                "GET /api/incidents/kpi/managers",
                "SUPERVISOR, EXECUTIVE.\n"
                "Ответ: KPI по менеджерам (назначения, решения, сроки, % в срок).",
                "403 — роль не подходит.\n200 [] — нет данных.",
            ),
            (
                "GET /api/incidents/overview",
                "SUPERVISOR, EXECUTIVE.\n"
                "Ответ: операционная панель — конвейер, «залипшие» инциденты, топ рисков, просрочки планов.",
                "403 — роль не подходит.",
            ),
            (
                "GET /api/v1/cases/my/stats",
                "MANAGER.\n"
                "Ответ: счётчики кейсов по статусам и среднее время закрытия.",
                "401 — не MANAGER.",
            ),
            (
                "GET /api/tasks/my/stats",
                "MANAGER.\n"
                "Ответ: сводка по задачам — всего, в работе, просрочено, сроки сегодня/завтра.",
                "401 — не MANAGER.",
            ),
        ],
    ),
]


USE_CASE_FIELD_LABELS = (
    "Наименование варианта использования",
    "Запрос",
    "Описание запроса с параметрами и заголовками",
    "Исключительные ситуации",
)


def add_use_case_request_table(
    doc: Document,
    use_case_name: str,
    request: str,
    description: str,
    exceptions: str,
) -> None:
    """Таблица 2×4: подписи в левом столбце, значения справа."""
    table = doc.add_table(rows=len(USE_CASE_FIELD_LABELS), cols=2)
    table.style = "Table Grid"
    values = (use_case_name, request, description, exceptions)

    for i, label in enumerate(USE_CASE_FIELD_LABELS):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        label_cell.text = label
        value_cell.text = values[i]
        for run in label_cell.paragraphs[0].runs:
            run.bold = True

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_use_case_api_section(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Варианты использования: запросы, ответы и исключения", level=1)
    doc.add_paragraph(
        "Описание HTTP-взаимодействий при реализации ключевых сценариев. "
        "Для каждого запроса — таблица: названия полей в левом столбце, значения в правом."
    ).paragraph_format.space_after = Pt(12)

    for use_case_name, rows in USE_CASE_API:
        doc.add_heading(use_case_name, level=2)
        doc.add_paragraph(f"Запросов в сценарии: {len(rows)}.")

        for request, description, exceptions in rows:
            add_use_case_request_table(
                doc, use_case_name, request, description, exceptions
            )

        doc.add_paragraph()


def add_table(doc: Document, rows: list[tuple[str, str, str, str, str]]) -> None:
    row_count = len(rows)
    doc.add_paragraph(f"Строк в таблице (без шапки): {row_count}")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("№", "Функция", "Метод / канал", "Роли", "Примечание")
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.add_paragraph()


def main() -> None:
    doc = Document()

    title = doc.add_heading(
        "Функциональные возможности cms-workflow-service",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Перечень реализованных возможностей микросервиса TrustFlow CMS Workflow "
        "с указанием ролей пользователей CMS Auth, которым доступна каждая функция. "
        "Источник: REST API (NestJS) и фоновые процессы по состоянию кода репозитория."
    )
    intro.paragraph_format.space_after = Pt(12)

    doc.add_heading("Роли пользователей", level=2)
    legend_table = doc.add_table(rows=1, cols=2)
    legend_table.style = "Table Grid"
    legend_table.rows[0].cells[0].text = "Роль"
    legend_table.rows[0].cells[1].text = "Описание"
    for cell in legend_table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for role, desc in ROLES_LEGEND:
        row = legend_table.add_row().cells
        row[0].text = role
        row[1].text = desc
    doc.add_paragraph()

    total = sum(len(s[2]) for s in SECTIONS)
    doc.add_paragraph(f"Всего функций в перечне: {total}.").paragraph_format.space_after = Pt(18)

    for section_title, section_note, rows in SECTIONS:
        doc.add_heading(section_title, level=1)
        if section_note:
            doc.add_paragraph(section_note)
        add_table(doc, rows)

    add_use_case_api_section(doc)

    output = Path(os.environ.get("DOCX_OUTPUT", OUTPUT))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Written: {output}")


if __name__ == "__main__":
    main()
