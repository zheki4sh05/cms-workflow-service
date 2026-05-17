"""Генерация Word-документа с описанием физической схемы PostgreSQL."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = Path(__file__).resolve().parent.parent / "physical-schema-postgresql.docx"

# (атрибут, тип, описание)
ENTITIES: list[dict] = [
    {
        "title": "Исходящее сообщение",
        "table": "outbox_messages",
        "description": (
            "Запись очереди Transactional Outbox для надёжной асинхронной обработки. "
            "После приёма события из Kafka сервис сохраняет payload в outbox_messages; "
            "фоновый OutboxProcessorService обрабатывает сообщения "
            "(например, topic incident_topic.received → создание инцидента через IncidentResolverService)."
        ),
        "notes": "Индекс: (status, createdAt). Внешних ключей на доменные таблицы нет.",
        "rows": [
            ("id", "uuid", "PK. UUID сообщения"),
            ("topic", "varchar(255)", "Логический тип события для обработчика (напр. incident_topic.received)"),
            ("payload", "jsonb", "JSON-тело события"),
            ("createdAt", "timestamptz", "Время постановки в outbox"),
            ("status", "varchar(20)", "pending | processed | failed"),
            ("processedAt", "timestamptz", "Время завершения обработки"),
            ("errorMessage", "text", "Текст ошибки при failed"),
        ],
    },
    {
        "title": "Инцидент",
        "table": "incident",
        "description": (
            "Корневая сущность workflow: агрегат комплаенс-события по объекту риска компании. "
            "Создаётся после разбора входящего Kafka-сообщения. При переходе в RESOLVED заполняется resolved_date."
        ),
        "rows": [
            ("id", "uuid", "PK. UUID инцидента"),
            ("companyId", "varchar(255)", "Идентификатор компании (tenant)"),
            ("integrationId", "integer", "ID интеграции мониторинга"),
            ("riskObjectId", "varchar(255)", "ID объекта риска (monitoring)"),
            ("departmentId", "varchar(255)", "ID подразделения (обогащение из monitoring)"),
            ("documentId", "varchar(255)", "ID связанного документа из события"),
            ("status", "varchar(20)", "OPEN, PARTLY_PROGRESS, IN_PROGRESS, RESOLVED"),
            ("resolved_date", "timestamp", "Момент перехода в RESOLVED"),
        ],
    },
    {
        "title": "Обнаружения (находка)",
        "table": "findings",
        "description": (
            "Отдельное срабатывание правила внутри инцидента. "
            "Одна находка может породить несколько кейсов; у каждого кейса ровно одна находка-родитель."
        ),
        "rows": [
            ("id", "uuid", "PK. UUID находки"),
            ("priority", "varchar(50)", "Приоритет срабатывания правила"),
            ("assignedUserId", "varchar(255)", "Предлагаемый ответственный"),
            ("rulesId", "uuid", "UUID правила (risk-сервис)"),
            ("detectedAt", "timestamptz", "Время обнаружения"),
            ("deadline", "timestamptz", "Крайний срок обработки"),
            ("details", "jsonb", "Детали срабатывания (result, found, details, …)"),
            ("incidentId", "uuid", "FK → incident.id"),
        ],
    },
    {
        "title": "Кейс",
        "table": "cases",
        "description": (
            "Единица работы исполнителя по конкретной находке в рамках инцидента. "
            "Через кейс доступны расследование, комментарии, вложения и план действий."
        ),
        "notes": "Уникальность: (incidentId, findingId, assignedUserId).",
        "rows": [
            ("id", "uuid", "PK. UUID кейса"),
            ("incidentId", "uuid", "FK → incident.id"),
            ("findingId", "uuid", "FK → findings.id"),
            ("assignedUserId", "varchar(255)", "Исполнитель; входит в uq_cases_incident_assignee"),
            ("status", "varchar(20)", "Этап workflow кейса (default OPEN)"),
        ],
    },
    {
        "title": "Расследование",
        "table": "investigations",
        "description": (
            "Материалы этапа расследования по кейсу. Связь с кейсом 1:1 — "
            "у кейса не более одного расследования."
        ),
        "rows": [
            ("id", "uuid", "PK. UUID расследования"),
            ("caseId", "uuid", "FK → cases.id (уникально)"),
            ("investigationNotes", "text", "Заметки по расследованию"),
            ("rootCause", "text", "Установленная первопричина"),
            ("requiresCorrectiveAction", "boolean", "Нужен ли план корректирующих действий"),
            ("createdAt", "timestamptz", "Время создания"),
            ("updatedAt", "timestamptz", "Время последнего изменения"),
        ],
    },
    {
        "title": "Комментарий к кейсу",
        "table": "case_comments",
        "description": "Текстовое сообщение участника в ленте обсуждения кейса.",
        "rows": [
            ("id", "uuid", "PK. UUID комментария"),
            ("caseId", "uuid", "FK → cases.id"),
            ("userId", "varchar(255)", "Автор (CMS Auth)"),
            ("comment", "text", "Текст комментария"),
            ("time", "timestamptz", "Время публикации"),
        ],
    },
    {
        "title": "Вложение к кейсу",
        "table": "case_attachments",
        "description": "Метаданные файла, прикреплённого к кейсу. Содержимое хранится в MinIO.",
        "rows": [
            ("id", "uuid", "PK. UUID записи"),
            ("caseId", "uuid", "FK → cases.id"),
            ("userId", "varchar(255)", "Кто загрузил файл"),
            ("fileId", "varchar(255)", "Ключ объекта в MinIO"),
            ("time", "timestamptz", "Время загрузки"),
            ("name", "varchar(512)", "Имя файла"),
            ("size", "integer", "Размер в байтах (default 0)"),
        ],
    },
    {
        "title": "План действий",
        "table": "action_plans",
        "description": (
            "План корректирующих мер по кейсу. Привязан к кейсу и инциденту. "
            "На один кейс — не более одного плана."
        ),
        "rows": [
            ("id", "uuid", "PK. UUID плана"),
            ("caseId", "uuid", "FK → cases.id (уникально)"),
            ("incidentId", "uuid", "FK → incident.id"),
            ("title", "varchar(500)", "Заголовок"),
            ("description", "text", "Описание мер"),
            ("comment", "text", "Комментарий при согласовании"),
            ("showTasks", "boolean", "Видимость задач (default false)"),
        ],
    },
    {
        "title": "Задача плана",
        "table": "action_plan_tasks",
        "description": "Конкретное действие в плане корректирующих мер.",
        "rows": [
            ("id", "uuid", "PK. UUID задачи"),
            ("actionPlanId", "uuid", "FK → action_plans.id"),
            ("title", "varchar(255)", "Название"),
            ("description", "text", "Описание задачи"),
            ("priority", "varchar(20)", "LOW, NORMAL, HIGH, CRITICAL"),
            ("dueDate", "timestamptz", "Срок исполнения"),
            ("status", "varchar(20)", "TODO, IN_PROGRESS, DONE (CHECK)"),
            ("evidenceDescriptionInprogress", "text", "Требования к доказательствам (in progress)"),
            ("evidenceDescriptionDone", "text", "Требования к доказательствам (done)"),
            ("completedAt", "timestamptz", "Время завершения"),
        ],
    },
    {
        "title": "Доказательство по задаче",
        "table": "action_plan_task_evidences",
        "description": "Файл-доказательство выполнения задачи. Объект в MinIO.",
        "rows": [
            ("id", "uuid", "PK. UUID записи"),
            ("taskId", "uuid", "FK → action_plan_tasks.id"),
            ("userId", "varchar(255)", "Кто прикрепил файл"),
            ("fileId", "uuid", "UUID объекта в MinIO"),
            ("name", "varchar(500)", "Имя файла"),
            ("time", "timestamptz", "Время прикрепления"),
        ],
    },
    {
        "title": "Верификация плана",
        "table": "verifications",
        "description": (
            "Результат проверки плана супервайзером. "
            "На один план — не более одной записи верификации."
        ),
        "rows": [
            ("id", "uuid", "PK. UUID верификации"),
            ("actionPlanId", "uuid", "FK → action_plans.id (уникально)"),
            ("verified", "boolean", "План принят (default false)"),
            ("assignedUserForVerification", "varchar(255)", "ID верификатора (Auth)"),
            ("assignedEmployeeForVerification", "varchar(255)", "ID сотрудника (company-info)"),
            ("comments", "text", "Комментарий супервайзера"),
        ],
    },
]


def add_entity(doc: Document, entity: dict) -> None:
    doc.add_heading(f"{entity['title']} ({entity['table']})", level=2)
    doc.add_paragraph(entity["description"])

    row_count = len(entity["rows"])
    count_para = doc.add_paragraph(f"Строк в таблице (без шапки): {row_count}")
    count_para.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Атрибут", "Тип", "Описание")
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True

    for row in entity["rows"]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    notes = entity.get("notes")
    if notes:
        p = doc.add_paragraph(notes)
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()


def main() -> None:
    doc = Document()
    title = doc.add_heading("Физическая схема PostgreSQL", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Микросервис cms-workflow-service. Итоговое состояние схемы по миграциям TypeORM. "
        "DDL-скрипт: database/schema/create-physical-schema.sql. "
        "Имена колонок совпадают с TypeORM (camelCase; исключение — incident.resolved_date)."
    )
    intro.paragraph_format.space_after = Pt(18)

    for entity in ENTITIES:
        add_entity(doc, entity)

    output = Path(os.environ.get("DOCX_OUTPUT", OUTPUT))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Written: {output}")


if __name__ == "__main__":
    main()
